# ✅ Standard Library
import os
import sys
import csv
import uuid
import tempfile
import logging
import urllib.parse
from datetime import datetime
from io import BytesIO

# ✅ Third-Party Libraries
import requests
from requests.exceptions import RequestException
from fastapi import FastAPI, Request, HTTPException, Depends
from fastapi.responses import (
    RedirectResponse,
    JSONResponse,
    HTMLResponse,
    StreamingResponse,
    Response,
)
from fastapi.middleware.cors import CORSMiddleware
import openpyxl
from PyPDF2 import PdfReader
from docx import Document




app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

logging.basicConfig(
    level=logging.INFO,
    format="%(message)s",
    handlers=[logging.StreamHandler(sys.stdout)]
)
logger = logging.getLogger("human-friendly-logger")

@app.middleware("http")
async def user_friendly_logger(request: Request, call_next):
    user_id = request.query_params.get("user_id", "unknown_user")
    ip = request.client.host if request.client else "unknown_ip"
    path = request.url.path
    method = request.method
    timestamp = datetime.utcnow().strftime("%b %d %Y %I:%M:%S %p")

    action_map = {
        "/oauth/login": "started login",
        "/oauth/callback": "completed login",
        "/generate_test_case": "requested test case generation",
        "/report_defect": "reported a defect",
        "/impact_analysis": "ran impact analysis",
        "/ticket": "created a Jira ticket",
        "/save-output": "saved generated output",
        "/export/docx": "exported DOCX file",
        "/export/csv": "exported CSV file",
        "/": "pinged home"
    }
    action = action_map.get(path, f"accessed {path}")

    logger.info(f"\n📥 REQUEST | [{timestamp}]")
    logger.info(f"User ID: {user_id}")
    logger.info(f"IP Address: {ip}")
    logger.info(f"Action: {action}")
    logger.info(f"Method: {method}")
    logger.info(f"Path: {path}")

    response: Response = await call_next(request)
    logger.info(f"📤 RESPONSE | Status Code: {response.status_code}")
    logger.info("-" * 60)
    return response

@app.get("/")
async def home():
    return {"message": "Welcome to QA GPT backend"}

CLIENT_ID = os.getenv("ATLASSIAN_CLIENT_ID")
CLIENT_SECRET = os.getenv("ATLASSIAN_CLIENT_SECRET")
REDIRECT_URI = os.getenv("OAUTH_REDIRECT_URI", "").strip()
AUTH_BASE_URL = "https://auth.atlassian.com/authorize"
TOKEN_URL = "https://auth.atlassian.com/oauth/token"
USER_API_URL = "https://api.atlassian.com/me"
RESOURCE_API = "https://api.atlassian.com/oauth/token/accessible-resources"
SCOPES = ["read:jira-work", "write:jira-work", "read:jira-user"]
user_tokens = {}

@app.get("/oauth/login")
def start_oauth():
    query = {
        "audience": "api.atlassian.com",
        "client_id": CLIENT_ID,
        "scope": " ".join(SCOPES),
        "redirect_uri": REDIRECT_URI,
        "state": "secureState123",
        "response_type": "code",
        "prompt": "consent"
    }
    url = f"{AUTH_BASE_URL}?{urllib.parse.urlencode(query)}"
    return RedirectResponse(url)

@app.get("/oauth/callback")
async def oauth_callback(request: Request):
    code = request.query_params.get("code")
    if not code:
        raise HTTPException(status_code=400, detail="Missing code")
    token_response = requests.post(TOKEN_URL, json={
        "grant_type": "authorization_code",
        "client_id": CLIENT_ID,
        "client_secret": CLIENT_SECRET,
        "code": code,
        "redirect_uri": REDIRECT_URI
    })
    if token_response.status_code != 200:
        return JSONResponse(status_code=token_response.status_code, content=token_response.json())
    tokens = token_response.json()
    access_token = tokens["access_token"]
    headers = {"Authorization": f"Bearer {access_token}"}
    user_info = requests.get(USER_API_URL, headers=headers).json()
    user_id = user_info.get("account_id", f"user_{uuid.uuid4()}")
    cloud_info = requests.get(RESOURCE_API, headers=headers).json()
    if not cloud_info:
        raise HTTPException(status_code=400, detail="No accessible Jira site found")
    cloud_id = cloud_info[0]["id"]
    base_url = f"https://api.atlassian.com/ex/jira/{cloud_id}"
    user_tokens[user_id] = {
        "access_token": access_token,
        "cloud_id": cloud_id,
        "base_url": base_url
    }
    logger.info(f"OAuth Success | User: {user_id}")
    return HTMLResponse(content=f"""
        <h2>✅ Jira OAuth Login Successful!</h2>
        <p>You can now return to ChatGPT and continue using the assistant.</p>
        <p><strong>Your Jira ID:</strong> <code>{user_id}</code></p>
        <p>📋 Please copy this ID and paste it into ChatGPT when asked.</p>
    """, status_code=200)

def get_auth_headers(request: Request):
    x_user_id = request.query_params.get("user_id")
    if not x_user_id:
        raise HTTPException(status_code=401, detail="Missing user_id. Please log in via /oauth/login.")
    if x_user_id not in user_tokens:
        raise HTTPException(status_code=401, detail="Session expired or user not authenticated.")
    data = user_tokens[x_user_id]
    if not data.get("project_key"):
        raise HTTPException(status_code=400, detail="Project key not set. Please call /set-project first.")
    return {
        "Authorization": f"Bearer {data['access_token']}",
        "Accept": "application/json",
        "Content-Type": "application/json"
    }, data["base_url"], data["project_key"]

@app.post("/set-project")
async def set_project(request: Request):
    x_user_id = request.query_params.get("user_id")
    if not x_user_id or x_user_id not in user_tokens:
        raise HTTPException(status_code=401, detail="Unauthorized or session expired")
    body = await request.json()
    project_key = body.get("project_key")
    if not project_key:
        raise HTTPException(status_code=400, detail="Missing 'project_key' in body")
    user_tokens[x_user_id]["project_key"] = project_key
    return {"message": f"Project key '{project_key}' set successfully"}

# Get Projects
@app.get("/projects")
async def get_projects(request: Request, auth_data=Depends(get_auth_headers)):
    headers, base_url, _ = auth_data
    try:
        res = requests.get(f"{base_url}/rest/api/3/project", headers=headers)
        res.raise_for_status()  # Raise error for 4xx/5xx responses
        return res.json()
    except RequestException as e:
        # Handles network issues, timeouts, invalid responses
        return JSONResponse(status_code=500, content={"error": f"Jira API request failed: {str(e)}"})
    except ValueError:
        # Handles JSON decode errors
        return JSONResponse(status_code=500, content={"error": "Invalid JSON response from Jira"})

# Get Tickets
@app.get("/ticket/{issue_key}")
async def fetch_ticket(issue_key: str, request: Request, auth_data=Depends(get_auth_headers)):
    headers, base_url, _ = auth_data

    # Get ticket info
    issue_url = f"{base_url}/rest/api/3/issue/{issue_key}"
    issue_res = requests.get(issue_url, headers=headers)
    if issue_res.status_code != 200:
        return JSONResponse(status_code=issue_res.status_code, content={"error": issue_res.text})

    issue = issue_res.json()
    description = issue["fields"].get("description", "")
    summary = issue["fields"].get("summary", "")

    # Get comments
    comments_url = f"{issue_url}/comment"
    comments_res = requests.get(comments_url, headers=headers)
    comments = comments_res.json().get("comments", []) if comments_res.status_code == 200 else []

    # Get attachments
    attachments = []
    for att in issue["fields"].get("attachment", []):
        filename = att.get("filename", "")
        content_url = att.get("content")
        if content_url:
            file_res = requests.get(content_url, headers=headers)
            if file_res.status_code == 200:
                text = extract_text_from_attachment(filename, file_res.content)
                attachments.append({
                    "filename": filename,
                    "content": text
                })

    return {
        "key": issue_key,
        "summary": summary,
        "description": description,
        "comments": comments,
        "attachments": attachments
    }


# Get attachements
def extract_text_from_attachment(filename, file_bytes):
    ext = os.path.splitext(filename)[-1].lower()

    try:
        if ext == ".txt":
            return file_bytes.decode("utf-8", errors="ignore")

        elif ext == ".pdf":
            reader = PdfReader(BytesIO(file_bytes))
            return "\n".join(page.extract_text() or "" for page in reader.pages)

        elif ext == ".docx":
            doc = Document(BytesIO(file_bytes))
            return "\n".join(p.text for p in doc.paragraphs)

        elif ext == ".csv":
            decoded = file_bytes.decode("utf-8", errors="ignore")
            lines = list(csv.reader(decoded.splitlines()))
            return "\n".join([", ".join(row) for row in lines])

        elif ext == ".xlsx":
            with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as tmp:
                tmp.write(file_bytes)
                tmp.flush()
                wb = openpyxl.load_workbook(tmp.name)
                content = []
                for sheet in wb.worksheets:
                    for row in sheet.iter_rows(values_only=True):
                        content.append(", ".join([str(cell) if cell is not None else "" for cell in row]))
                os.unlink(tmp.name)
                return "\n".join(content)

        else:
            return f"[Unsupported file type: {ext}]"

    except Exception as e:
        return f"[ERROR reading {filename}: {str(e)}]"

# Create Tickets
@app.post("/ticket")
async def create_ticket(request: Request, auth_data=Depends(get_auth_headers)):
    headers, base_url, project_key = auth_data
    data = await request.json()
    payload = {
        "fields": {
            "project": {"key": project_key},
            "summary": data.get("summary"),
            "description": {
                "type": "doc", "version": 1,
                "content": [{"type": "paragraph", "content": [{"type": "text", "text": data.get("description", "")}]}]
            },
            "issuetype": {"name": data.get("issue_type", "Bug")}
        }
    }
    try:
        res = requests.post(f"{base_url}/rest/api/3/issue", headers=headers, json=payload)
        res.raise_for_status()
        return {"message": "Ticket created", "key": res.json().get("key")}
    except requests.exceptions.RequestException as e:
        return JSONResponse(status_code=500, content={"error": f"Failed to create ticket: {str(e)}"})
    except ValueError:
        return JSONResponse(status_code=500, content={"error": "Invalid response from Jira"})

# Update Tickets
@app.patch("/ticket/{issue_key}")
async def update_ticket(issue_key: str, request: Request, auth_data=Depends(get_auth_headers)):
    headers, base_url, _ = auth_data
    try:
        data = await request.json()
        update_fields = {}

        if "summary" in data:
            update_fields["summary"] = data["summary"]
        if "description" in data:
            update_fields["description"] = {
                "type": "doc", "version": 1,
                "content": [{"type": "paragraph", "content": [{"type": "text", "text": data["description"]}]}]
            }

        payload = {"fields": update_fields}
        res = requests.put(f"{base_url}/rest/api/3/issue/{issue_key}", headers=headers, json=payload)
        res.raise_for_status()

        return {"message": "Ticket updated successfully"} if res.status_code == 204 else res.json()

    except RequestException as e:
        return JSONResponse(status_code=500, content={"error": f"Failed to update ticket {issue_key}: {str(e)}"})
    except ValueError:
        return JSONResponse(status_code=500, content={"error": "Invalid response format from Jira"})

# Get Ticket comments
@app.get("/ticket/{issue_key}/comments")
async def get_comments(issue_key: str, request: Request, auth_data=Depends(get_auth_headers)):
    headers, base_url, _ = auth_data
    try:
        res = requests.get(f"{base_url}/rest/api/3/issue/{issue_key}/comment", headers=headers)
        res.raise_for_status()
        return res.json()
    except requests.exceptions.RequestException as e:
        return JSONResponse(status_code=500, content={"error": f"Failed to fetch comments: {str(e)}"})
    except ValueError:
        return JSONResponse(status_code=500, content={"error": "Invalid JSON from Jira"})

# Add comments in a Ticket
@app.post("/ticket/{issue_key}/comments")
async def add_comment(issue_key: str, request: Request, auth_data=Depends(get_auth_headers)):
    headers, base_url, _ = auth_data
    data = await request.json()
    payload = {
        "body": {
            "type": "doc", "version": 1,
            "content": [{"type": "paragraph", "content": [{"type": "text", "text": data.get("body", "")}]}]
        }
    }
    try:
        res = requests.post(f"{base_url}/rest/api/3/issue/{issue_key}/comment", headers=headers, json=payload)
        res.raise_for_status()
        return res.json()
    except requests.exceptions.RequestException as e:
        return JSONResponse(status_code=500, content={"error": f"Failed to add comment: {str(e)}"})
    except ValueError:
        return JSONResponse(status_code=500, content={"error": "Invalid JSON response from Jira"})

# Update Ticket Comments
@app.patch("/ticket/{issue_key}/comments/{comment_id}")
async def update_comment(issue_key: str, comment_id: str, request: Request, auth_data=Depends(get_auth_headers)):
    headers, base_url, _ = auth_data
    data = await request.json()
    payload = {
        "body": {
            "type": "doc", "version": 1,
            "content": [{"type": "paragraph", "content": [{"type": "text", "text": data.get("body", "")}]}]
        }
    }
    try:
        res = requests.put(f"{base_url}/rest/api/3/issue/{issue_key}/comment/{comment_id}", headers=headers, json=payload)
        res.raise_for_status()
        return {"message": "Comment updated"}
    except requests.exceptions.RequestException as e:
        return JSONResponse(status_code=500, content={"error": f"Failed to update comment: {str(e)}"})

# Querry filetrs for tickets
def jql_search(jql: str, headers, base_url, project_key=None, source_ticket_key: str = None):
    try:
        full_jql = f'project = "{project_key}" AND {jql}' if project_key else jql
        res = requests.get(f"{base_url}/rest/api/3/search", headers=headers, params={"jql": full_jql, "maxResults": 100})
        res.raise_for_status()

        issues = []
        for i in res.json().get("issues", []):
            key = i["key"]
            if source_ticket_key and key == source_ticket_key:
                continue
            issue_type = i["fields"].get("issuetype", {}).get("name", "").lower()
            description = i["fields"].get("description", "")

            # Optional filter logic
            if issue_type in ["epic", "parent"]:
                if not description:
                    continue
                text = description if isinstance(description, str) else str(description)
                logic_keywords = ["verify", "check", "validate", "flow", "test", "should"]
                if not any(word in text.lower() for word in logic_keywords):
                    continue

            issues.append({"key": key, "summary": i["fields"]["summary"], "description": description})
        return issues

    except requests.exceptions.RequestException as e:
        return JSONResponse(status_code=500, content={"error": f"Jira JQL search failed: {str(e)}"})
    except ValueError:
        return JSONResponse(status_code=500, content={"error": "Invalid JSON returned by Jira"})

@app.get("/impact/label/{label}")
async def get_impact_by_label(label: str, request: Request, auth_data=Depends(get_auth_headers)):
    headers, base_url, project_key = auth_data
    source_ticket_key = request.query_params.get("source_ticket")
    return jql_search(f'labels = "{label}"', headers, base_url, project_key, source_ticket_key)

@app.get("/impact/component/{component}")
async def get_impact_by_component(component: str, request: Request, auth_data=Depends(get_auth_headers)):
    headers, base_url, project_key = auth_data
    source_ticket_key = request.query_params.get("source_ticket")
    return jql_search(f'component = "{component}"', headers, base_url, project_key, source_ticket_key)

@app.get("/impact/module/{keyword}")
async def get_impact_by_module(keyword: str, request: Request, auth_data=Depends(get_auth_headers)):
    headers, base_url, project_key = auth_data
    source_ticket_key = request.query_params.get("source_ticket")
    return jql_search(f'summary ~ "{keyword}"', headers, base_url, project_key, source_ticket_key)

@app.get("/tickets/sprint/{sprint_id}")
async def get_tickets_by_sprint(sprint_id: int, request: Request, auth_data=Depends(get_auth_headers)):
    headers, base_url, project_key = auth_data
    return jql_search(f"sprint = {sprint_id}", headers, base_url, project_key=project_key)

@app.get("/tickets/priority/{priority}")
async def get_tickets_by_priority(priority: str, request: Request, auth_data=Depends(get_auth_headers)):
    headers, base_url, project_key = auth_data
    return jql_search(f'priority = "{priority}"', headers, base_url, project_key=project_key)


@app.post("/generate-docx")
async def generate_docx(request: Request):
    data = await request.json()
    text = data.get("output", "")
    if not text:
        return JSONResponse(status_code=400, content={"error": "Missing 'output' field."})

    try:
        doc = Document()
        for line in text.split('\n'):
            doc.add_paragraph(line)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=test-output.docx"}
        )
    except Exception as e:
        return JSONResponse(status_code=500, content={"error": f"Failed to generate DOCX: {str(e)}"})


@app.post("/generate-csv")
async def generate_csv(request: Request):
    data = await request.json()
    text = data.get("output", "")
    if not text:
        return JSONResponse(status_code=400, content={"error": "Missing 'output' field."})

    try:
        buffer = StringIO()
        writer = csv.writer(buffer)
        for line in text.split('\n'):
            writer.writerow([line])
        buffer.seek(0)
        byte_stream = BytesIO(buffer.getvalue().encode())

        return StreamingResponse(
            byte_stream,
            media_type="text/csv",
            headers={"Content-Disposition": "attachment; filename=test-output.csv"}
        )
    except Exception as e:
        return JSONResponse(status_code=500, content={"error": f"Failed to generate CSV: {str(e)}"})

# Login redirect to auth
@app.get("/login")
def legacy_login_redirect():
    return RedirectResponse(url="/oauth/login")

# App Health
@app.get("/health")
async def health_check():
    return {"status": "ok"}
